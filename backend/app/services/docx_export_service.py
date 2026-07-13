"""DOCX 报告导出服务 — 基于 python-docx 导出 7 段式安全分析报告."""

import json
import logging
from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.database import get_connection
from app.models.host import Host
from app.models.ai_analysis import AiAnalysisReport

logger = logging.getLogger(__name__)

# ── 全局样式配置 ──
TITLE_COLOR = RGBColor(0x1A, 0x3C, 0x6E)
SECTION_COLOR = RGBColor(0x2E, 0x5E, 0x8E)
SUBSECTION_COLOR = RGBColor(0x34, 0x49, 0x5E)
HIGHLIGHT_COLOR = RGBColor(0xC0, 0x39, 0x2B)
NORMAL_FONT_SIZE = Pt(11)
SMALL_FONT_SIZE = Pt(9)


class DocxExportService:
    """DOCX 报告导出服务."""

    @staticmethod
    def export(report_id: int) -> Optional[bytes]:
        """导出 incident_report 为 DOCX 字节流.

        Args:
            report_id: incident_reports 表的主键 ID.

        Returns:
            DOCX 文件的字节流，失败返回 None.
        """
        # 1. 获取报告数据
        with get_connection() as conn:
            row = conn.execute(
                "SELECT * FROM incident_reports WHERE id = ?", (report_id,),
            ).fetchone()
        if not row:
            logger.error("incident_report %d not found", report_id)
            return None
        report = dict(row)

        # 2. 获取关联的 AI 报告
        ai_report_id = report.get("ai_report_id")
        ai_report: Optional[dict] = None
        if ai_report_id:
            ai_report = AiAnalysisReport.get_by_id(ai_report_id)

        # 3. 获取主机信息
        host = Host.get_by_id(report.get("host_id", 0)) if report.get("host_id") else None

        # 4. 解析 JSON 字段
        impact_scope = DocxExportService._safe_parse_json(report.get("impact_scope", "[]"))
        timeline_json = DocxExportService._safe_parse_json(report.get("timeline_json", "[]"))
        mitre_cover = DocxExportService._safe_parse_json(report.get("mitre_cover", "[]"))
        recommendations_raw = DocxExportService._safe_parse_json(report.get("recommendations", "{}"))

        doc = Document()

        # ── 设置默认字体 ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = NORMAL_FONT_SIZE

        # ── 封面标题 ──
        title = doc.add_heading(report.get("title", "安全分析报告"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = TITLE_COLOR

        # 报告元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"报告类型: {report.get('report_type', 'N/A')}  |  ")
        meta.add_run(f"状态: {report.get('status', 'N/A')}  |  ")
        meta.add_run(f"生成时间: {report.get('created_at', '')}")
        for run in meta.runs:
            run.font.size = SMALL_FONT_SIZE
            run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

        doc.add_paragraph()  # 空行

        # ── 段落1: 报告概要 ──
        DocxExportService._add_section(doc, "一、报告概要", SECTION_COLOR)
        summary = report.get("summary", "") or "暂无概要"
        doc.add_paragraph(summary)

        # ── 段落2: 主机信息 ──
        DocxExportService._add_section(doc, "二、主机信息", SECTION_COLOR)
        if host:
            hostname = host.get("hostname", "N/A")
            ip = host.get("ip_address", "N/A")
            os_info = host.get("os_type", "N/A")
            os_ver = host.get("os_version", "")
            host_text = f"主机名: {hostname}  |  IP: {ip}  |  操作系统: {os_info}"
            if os_ver:
                host_text += f" ({os_ver})"
            doc.add_paragraph(host_text)
        else:
            doc.add_paragraph("主机信息: N/A")

        # ── 段落3: 风险分析 ──
        DocxExportService._add_section(doc, "三、风险分析", SECTION_COLOR)
        risk_score = report.get("risk_score", 0)
        doc.add_paragraph(f"风险评分: {risk_score}/100")

        confidence_raw = report.get("confidence_metadata")
        if confidence_raw:
            try:
                confidence = json.loads(confidence_raw) if isinstance(confidence_raw, str) else confidence_raw
                if isinstance(confidence, dict) and confidence.get("summary") is not None:
                    doc.add_paragraph(f"概要置信度: {confidence.get('summary', 'N/A')}%")
            except (json.JSONDecodeError, TypeError):
                pass

        # ── 段落4: 威胁分析（影响范围） ──
        DocxExportService._add_section(doc, "四、威胁分析", SECTION_COLOR)
        if isinstance(impact_scope, list) and impact_scope:
            DocxExportService._add_subsection(doc, "影响系统")
            for item in impact_scope:
                if isinstance(item, dict):
                    sys_name = item.get("system", item.get("name", str(item)))
                    sys_desc = item.get("description", "")
                    line = f"• {sys_name}"
                    if sys_desc:
                        line += f": {sys_desc}"
                    doc.add_paragraph(line, style="List Bullet")
                else:
                    doc.add_paragraph(f"• {item}", style="List Bullet")
        else:
            doc.add_paragraph("暂无影响范围数据")

        # ATT&CK 覆盖
        if isinstance(mitre_cover, list) and mitre_cover:
            DocxExportService._add_subsection(doc, "ATT&CK 覆盖")
            for tech in mitre_cover:
                if isinstance(tech, dict):
                    tid = tech.get("id", tech.get("technique_id", ""))
                    tname = tech.get("name", tech.get("technique_name", ""))
                    doc.add_paragraph(f"• {tid} {tname}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {tech}", style="List Bullet")

        # ── 段落5: 时间线 ──
        DocxExportService._add_section(doc, "五、时间线", SECTION_COLOR)
        if isinstance(timeline_json, list) and timeline_json:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "事件"
            hdr[2].text = "来源"

            for event in timeline_json:
                if isinstance(event, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(event.get("timestamp", event.get("time", "")))
                    row_cells[1].text = str(event.get("description", event.get("event", "")))
                    row_cells[2].text = str(event.get("source", ""))
        else:
            doc.add_paragraph("暂无时间线数据")

        # ── 段落6: 证据 ──
        DocxExportService._add_section(doc, "六、证据", SECTION_COLOR)
        evidence = report.get("evidence", "") or "暂无证据"
        # 移除证据链接标记 [🔗](...)
        import re
        evidence_clean = re.sub(r'\[🔗\]\([^)]+\)\s*', '', evidence)
        doc.add_paragraph(evidence_clean)

        # ── 段落7: 处置建议 ──
        DocxExportService._add_section(doc, "七、处置建议", SECTION_COLOR)
        if isinstance(recommendations_raw, dict):
            items = recommendations_raw.get("items", recommendations_raw.get("recommendations", []))
            if isinstance(items, list) and items:
                for idx, rec in enumerate(items, 1):
                    if isinstance(rec, dict):
                        title_text = rec.get("title", rec.get("action", f"建议 {idx}"))
                        desc = rec.get("description", rec.get("detail", ""))
                        priority = rec.get("priority", "")
                        status = rec.get("status", "")

                        p = doc.add_paragraph()
                        run = p.add_run(f"{idx}. {title_text}")
                        run.bold = True
                        if priority:
                            p.add_run(f"  [{priority}]")

                        if desc:
                            doc.add_paragraph(f"   {desc}")
                        if status:
                            doc.add_paragraph(f"  状态: {status}")
                    else:
                        doc.add_paragraph(f"{idx}. {rec}")
            else:
                doc.add_paragraph("暂无处置建议")
        else:
            doc.add_paragraph("暂无处置建议")

        # ── 页脚 ──
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("— 报告完 —")
        run.font.size = SMALL_FONT_SIZE
        run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

        # ── 保存到字节流 ──
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    @staticmethod
    def _add_section(doc: Document, title: str, color: RGBColor) -> None:
        """添加章节标题."""
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.color.rgb = color

    @staticmethod
    def _add_subsection(doc: Document, title: str) -> None:
        """添加子章节标题."""
        h = doc.add_heading(title, level=2)
        for run in h.runs:
            run.font.color.rgb = SUBSECTION_COLOR

    @staticmethod
    def _safe_parse_json(val: Any) -> Any:
        """安全解析 JSON 字符串."""
        if val is None:
            return None
        if isinstance(val, (dict, list)):
            return val
        if isinstance(val, str):
            try:
                return json.loads(val)
            except (json.JSONDecodeError, TypeError):
                return val
        return val
